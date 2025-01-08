from docx import Document

def main():
    print("=== Программа для создания Word-документа ===")
    
   
    user_text = input("Введите текст, который вы хотите сохранить в Word-документ: ")
    
    doc = Document()
    doc.add_heading("Документ пользователя", level=1)
    doc.add_paragraph(user_text)
    
    file_name = "output.docx"
    doc.save(file_name)
    
    print(f"Текст успешно сохранён в файле: {file_name}")

if __name__ == "__main__":
    main()